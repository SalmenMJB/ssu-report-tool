"""
Tests pour le générateur de rapports Word SSU.
"""

import os
import tempfile

import pytest
from docx import Document

from app.report_generator.document_builder import ReportBuilder
from app.report_generator.styles import COLOR_BLUE, FONT_NAME, MARGIN_CM


class TestReportBuilder:
    """Tests unitaires pour la classe ReportBuilder."""

    def _make_builder(self, tmp_path) -> ReportBuilder:
        output = os.path.join(str(tmp_path), "test_report.docx")
        return ReportBuilder(output)

    def test_builder_creates_document(self, tmp_path):
        """ReportBuilder crée un document Word valide."""
        builder = self._make_builder(tmp_path)
        builder.save()
        output = os.path.join(str(tmp_path), "test_report.docx")
        assert os.path.isfile(output)

    def test_margins_configured(self, tmp_path):
        """Le document utilise les marges définies dans styles.py (±1 % de tolérance)."""
        builder = self._make_builder(tmp_path)
        section = builder.document.sections[0]
        # Les marges EMU peuvent différer légèrement à cause des arrondis lors de
        # la conversion cm → EMU → cm. On accepte une tolérance de ±1 %.
        tolerance = int(MARGIN_CM * 0.01)
        assert abs(section.left_margin - MARGIN_CM) <= tolerance
        assert abs(section.right_margin - MARGIN_CM) <= tolerance
        assert abs(section.top_margin - MARGIN_CM) <= tolerance
        assert abs(section.bottom_margin - MARGIN_CM) <= tolerance

    def test_heading1_color(self, tmp_path):
        """Le Heading 1 doit être en bleu ciel UA."""
        builder = self._make_builder(tmp_path)
        h1 = builder.document.styles["Heading 1"]
        assert h1.font.color.rgb == COLOR_BLUE

    def test_add_title_page(self, tmp_path):
        """add_title_page insère le titre et le sous-titre."""
        builder = self._make_builder(tmp_path)
        builder.add_title_page("Rapport d'activité", "2024 – 2025")
        builder.save()

        doc = Document(os.path.join(str(tmp_path), "test_report.docx"))
        all_text = " ".join(p.text for p in doc.paragraphs)
        assert "RAPPORT D'ACTIVITÉ" in all_text.upper()
        assert "2024" in all_text

    def test_add_chapter(self, tmp_path):
        """add_chapter insère un Heading 1 avec le bon titre."""
        builder = self._make_builder(tmp_path)
        builder.add_chapter("Médecine générale", "Texte intro")
        builder.save()

        doc = Document(os.path.join(str(tmp_path), "test_report.docx"))
        headings = [p.text for p in doc.paragraphs if p.style.name == "Heading 1"]
        assert "Médecine générale" in headings

    def test_add_section(self, tmp_path):
        """add_section insère un Heading 2."""
        builder = self._make_builder(tmp_path)
        builder.add_section("Activité médicale", "Texte section")
        builder.save()

        doc = Document(os.path.join(str(tmp_path), "test_report.docx"))
        headings = [p.text for p in doc.paragraphs if p.style.name == "Heading 2"]
        assert "Activité médicale" in headings

    def test_add_key_indicators_table(self, tmp_path):
        """add_key_indicators_table crée un tableau avec des données."""
        builder = self._make_builder(tmp_path)
        indicators = {"Total consultations": 1500, "Étudiants uniques": 800}
        builder.add_key_indicators_table(indicators)
        builder.save()

        doc = Document(os.path.join(str(tmp_path), "test_report.docx"))
        assert len(doc.tables) >= 1

        # Vérifier que les valeurs sont dans le tableau
        table = doc.tables[0]
        all_table_text = " ".join(
            cell.text for row in table.rows for cell in row.cells
        )
        assert "Total consultations" in all_table_text
        assert "1500" in all_table_text

    def test_add_image_missing_file(self, tmp_path):
        """add_image insère un placeholder si le fichier n'existe pas."""
        builder = self._make_builder(tmp_path)
        builder.add_image("/nonexistent/chart.png", caption="Test")
        builder.save()

        doc = Document(os.path.join(str(tmp_path), "test_report.docx"))
        all_text = " ".join(p.text for p in doc.paragraphs)
        assert "Graphique manquant" in all_text

    def test_add_image_existing_file(self, tmp_path):
        """add_image insère l'image si le fichier existe."""
        # Créer une image PNG minimale avec Pillow
        from PIL import Image
        img_path = os.path.join(str(tmp_path), "test_chart.png")
        img = Image.new("RGB", (100, 100), color=(0, 100, 200))
        img.save(img_path)

        builder = self._make_builder(tmp_path)
        builder.add_image(img_path, caption="Graphique test")
        builder.save()

        doc = Document(os.path.join(str(tmp_path), "test_report.docx"))
        assert len(doc.inline_shapes) >= 1

    def test_add_editable_comment_zone(self, tmp_path):
        """add_editable_comment_zone insère le texte de remplacement."""
        builder = self._make_builder(tmp_path)
        builder.add_editable_comment_zone()
        builder.save()

        doc = Document(os.path.join(str(tmp_path), "test_report.docx"))
        all_text = " ".join(p.text for p in doc.paragraphs)
        assert "À compléter par l'équipe SSU" in all_text

    def test_add_table_of_contents(self, tmp_path):
        """add_table_of_contents insère le titre 'Sommaire'."""
        builder = self._make_builder(tmp_path)
        builder.add_table_of_contents()
        builder.save()

        doc = Document(os.path.join(str(tmp_path), "test_report.docx"))
        headings = [p.text for p in doc.paragraphs if p.style.name == "Heading 1"]
        assert "Sommaire" in headings

    def test_empty_indicators_table_skipped(self, tmp_path):
        """add_key_indicators_table n'ajoute rien pour un dict vide."""
        builder = self._make_builder(tmp_path)
        builder.add_key_indicators_table({})
        builder.save()

        doc = Document(os.path.join(str(tmp_path), "test_report.docx"))
        assert len(doc.tables) == 0

    def test_full_report_structure(self, tmp_path):
        """Le rapport complet contient tous les chapitres requis."""
        from app.report_generator.chapters.intro import build_intro_chapter
        from app.report_generator.chapters.effectifs import build_effectifs_chapter
        from app.report_generator.chapters.medecine import build_medecine_chapter
        from app.report_generator.chapters.ide import build_ide_chapter
        from app.report_generator.chapters.consommables import build_consommables_chapter
        from app.report_generator.chapters.pssm import build_pssm_chapter
        from app.report_generator.chapters.dspe import (
            build_psy_chapter,
            build_psychiatrie_chapter,
            build_sante_mentale_chapter,
        )
        from app.report_generator.chapters.css import build_css_chapter
        from app.report_generator.chapters.partenariats import (
            build_dietetique_chapter,
            build_partenariats_chapter,
        )

        builder = self._make_builder(tmp_path)
        builder.add_title_page()
        builder.add_table_of_contents()

        build_intro_chapter(builder, {})
        build_effectifs_chapter(builder, None, None, {}, {})
        build_medecine_chapter(builder, None, {})
        build_ide_chapter(builder, {})
        build_consommables_chapter(builder, {})
        build_psy_chapter(builder, {})
        build_psychiatrie_chapter(builder, {})
        build_sante_mentale_chapter(builder)
        build_pssm_chapter(builder, {})
        build_css_chapter(builder, {})
        build_dietetique_chapter(builder)
        build_partenariats_chapter(builder)

        builder.save()

        doc = Document(os.path.join(str(tmp_path), "test_report.docx"))
        h1_titles = [p.text for p in doc.paragraphs if p.style.name == "Heading 1"]

        expected_chapters = [
            "Présentation générale du service",
            "Pôle administratif",
            "Médecine générale",
            "Service infirmier",
            "Éducation à la santé",
            "Psychologie",
            "Psychiatrie",
            "Santé mentale et bien-être",
            "Centre de santé sexuelle",
            "Diététique et Nutrition",
            "Partenariats & Perspectives",
        ]

        for chapter in expected_chapters:
            assert chapter in h1_titles, f"Chapitre manquant : {chapter}"
