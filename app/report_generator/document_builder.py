"""
Classe principale de construction du rapport Word SSU.

Usage :
    from app.report_generator.document_builder import ReportBuilder

    builder = ReportBuilder("output/rapport_ssu_2024_2025.docx")
    builder.add_title_page("Rapport d'activité", "2024 – 2025")
    builder.add_table_of_contents()
    builder.add_chapter("Présentation générale du service", "Texte intro...")
    builder.add_image("output/charts/recap_consultations.png", caption="Récapitulatif")
    builder.save()
"""

from __future__ import annotations

import os
from typing import List, Optional, Union

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor

from app.report_generator.styles import (
    ARROW_LEVEL1,
    ARROW_LEVEL2,
    COLOR_BLACK,
    COLOR_BLUE,
    COLOR_BLUE_LIGHT,
    COLOR_BLUE_PALE,
    COLOR_GRAY_LIGHT,
    COLOR_TABLE_HEADER,
    COLOR_TABLE_ROW_EVEN,
    COLOR_TABLE_ROW_ODD,
    COLOR_UA_BLUE,
    COLOR_WHITE,
    FONT_NAME,
    FONT_SIZE_BODY,
    FONT_SIZE_CAPTION,
    FONT_SIZE_H1,
    FONT_SIZE_H2,
    FONT_SIZE_H3,
    FONT_SIZE_SUBTITLE,
    FONT_SIZE_TITLE,
    IMAGE_WIDTH_FULL,
    MARGIN_CM,
)


# ---------------------------------------------------------------------------
# Helpers XML bas niveau
# ---------------------------------------------------------------------------

def _set_cell_background(cell, color: RGBColor) -> None:
    """Définit la couleur de fond d'une cellule de tableau."""
    hex_color = str(color)  # RGBColor.__str__ retourne le code hexadécimal, ex. "00A8E1"
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _set_run_font(run, name: str, size: Pt, bold: bool = False,
                  color: Optional[RGBColor] = None) -> None:
    run.font.name = name
    run.font.size = size
    run.font.bold = bold
    if color is not None:
        run.font.color.rgb = color


def _set_cell_border(cell, border_color_hex: str = "00A8E1") -> None:
    """Définit une bordure colorée sur une cellule."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for border_tag in ("w:top", "w:left", "w:bottom", "w:right"):
        border = OxmlElement(border_tag)
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "12")
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), border_color_hex)
        tcPr.append(border)


# ---------------------------------------------------------------------------
# Classe principale
# ---------------------------------------------------------------------------

class ReportBuilder:
    """Construit un rapport Word structuré au format SSU / Université d'Angers."""

    def __init__(self, output_path: str) -> None:
        self.output_path = output_path
        self.document = Document()
        self._configure_document()

    # ------------------------------------------------------------------
    # Configuration initiale
    # ------------------------------------------------------------------

    def _configure_document(self) -> None:
        """Applique les marges et configure les styles de base."""
        for section in self.document.sections:
            section.top_margin = MARGIN_CM
            section.bottom_margin = MARGIN_CM
            section.left_margin = MARGIN_CM
            section.right_margin = MARGIN_CM

        # Style Normal
        normal = self.document.styles["Normal"]
        normal.font.name = FONT_NAME
        normal.font.size = FONT_SIZE_BODY
        normal.paragraph_format.space_before = Pt(4)
        normal.paragraph_format.space_after = Pt(8)
        normal.paragraph_format.line_spacing = 1.15 # Interligne 1.15 (plus flexible que Pt(14))

        # Forcer les titres à avoir un interligne automatique pour ne pas être coupés
        for style_name in ["Heading 1", "Heading 2", "Heading 3"]:
            style = self.document.styles[style_name]
            style.paragraph_format.line_spacing = None # Reset l'interligne (Auto)
            style.paragraph_format.line_spacing_rule = None

        # Heading 1 — bleu ciel UA, grand, sans numérotation visible
        h1 = self.document.styles["Heading 1"]
        h1.font.name = FONT_NAME
        h1.font.size = FONT_SIZE_H1
        h1.font.bold = True
        h1.font.color.rgb = COLOR_BLUE
        h1.paragraph_format.space_before = Pt(18)
        h1.paragraph_format.space_after = Pt(12)

        # Heading 2 — bleu ciel UA
        h2 = self.document.styles["Heading 2"]
        h2.font.name = FONT_NAME
        h2.font.size = FONT_SIZE_H2
        h2.font.bold = True
        h2.font.color.rgb = COLOR_BLUE_LIGHT
        h2.paragraph_format.space_before = Pt(14)
        h2.paragraph_format.space_after = Pt(8)

        # Heading 3 — bleu ciel UA, légèrement plus petit
        h3 = self.document.styles["Heading 3"]
        h3.font.name = FONT_NAME
        h3.font.size = FONT_SIZE_H3
        h3.font.bold = True
        h3.font.color.rgb = COLOR_BLUE_LIGHT
        h3.paragraph_format.space_before = Pt(10)
        h3.paragraph_format.space_after = Pt(6)

        # Numéro de page dans le pied de page
        self._add_page_numbers()

    def _add_page_numbers(self) -> None:
        """Ajoute la numérotation des pages dans le pied de page."""
        section = self.document.sections[0]
        footer = section.footer
        footer_para = footer.paragraphs[0]
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer_para.clear()

        run = footer_para.add_run()
        run.font.name = FONT_NAME
        run.font.size = FONT_SIZE_CAPTION

        # Champ PAGE
        fld_char1 = OxmlElement("w:fldChar")
        fld_char1.set(qn("w:fldCharType"), "begin")

        instr_text = OxmlElement("w:instrText")
        instr_text.set(qn("xml:space"), "preserve")
        instr_text.text = "PAGE"

        fld_char2 = OxmlElement("w:fldChar")
        fld_char2.set(qn("w:fldCharType"), "end")

        run._r.append(fld_char1)
        run._r.append(instr_text)
        run._r.append(fld_char2)

    # ------------------------------------------------------------------
    # Page de titre
    # ------------------------------------------------------------------

    def add_title_page(
        self,
        title: str = "Rapport d'activité",
        subtitle: str = "2024 – 2025",
        logo_path: Optional[str] = None,
    ) -> None:
        """Génère la page de titre avec design minimaliste : logo, titre et sous-titre."""

        # Logo — on n'insère que si le fichier existe (taille réduite : 4,5 cm)
        if logo_path and os.path.isfile(logo_path):
            logo_para = self.document.add_paragraph()
            logo_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = logo_para.add_run()
            run.add_picture(logo_path, width=Cm(4.5))

        # Espace avant le titre
        for _ in range(4):
            self.document.add_paragraph()

        # Titre principal
        title_para = self.document.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title_para.add_run(title.upper())
        _set_run_font(run, FONT_NAME, FONT_SIZE_TITLE, bold=True, color=COLOR_UA_BLUE)

        # Ligne décorative bleue sous le titre
        separator_para = self.document.add_paragraph()
        separator_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sep_run = separator_para.add_run("─" * 30)
        sep_run.font.color.rgb = COLOR_UA_BLUE
        sep_run.font.size = Pt(12)

        # Sous-titre (année)
        subtitle_para = self.document.add_paragraph()
        subtitle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = subtitle_para.add_run(subtitle)
        _set_run_font(run, FONT_NAME, FONT_SIZE_SUBTITLE, bold=False, color=COLOR_BLACK)

        # Espace + Service
        self.document.add_paragraph()
        service_para = self.document.add_paragraph()
        service_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = service_para.add_run("Service de Santé Universitaire")
        _set_run_font(run, FONT_NAME, FONT_SIZE_H2, bold=True, color=COLOR_UA_BLUE)

        univ_para = self.document.add_paragraph()
        univ_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = univ_para.add_run("Université d'Angers")
        _set_run_font(run, FONT_NAME, FONT_SIZE_H3, bold=False, color=COLOR_BLACK)

        # Saut de page (page blanche entre couverture et sommaire)
        self.document.add_page_break()
        self.document.add_paragraph()
        self.document.add_page_break()

    # ------------------------------------------------------------------
    # Table des matières
    # ------------------------------------------------------------------

    def add_table_of_contents(self, title: str = "Sommaire") -> None:
        """
        Insère un champ TABLE DES MATIÈRES Word.
        La table est générée par Word/LibreOffice à l'ouverture du document.
        """
        toc_heading = self.document.add_heading(title, level=1)
        toc_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

        para = self.document.add_paragraph()
        run = para.add_run()

        fld_char1 = OxmlElement("w:fldChar")
        fld_char1.set(qn("w:fldCharType"), "begin")
        fld_char1.set(qn("w:dirty"), "true")

        instr_text = OxmlElement("w:instrText")
        instr_text.set(qn("xml:space"), "preserve")
        instr_text.text = 'TOC \\o "1-3" \\h \\z \\u'

        fld_char2 = OxmlElement("w:fldChar")
        fld_char2.set(qn("w:fldCharType"), "separate")

        fld_text = OxmlElement("w:t")
        fld_text.text = (
            "Cliquez ici, puis sur F9 (Windows) ou Cmd+A+F9 (Mac) "
            "pour mettre à jour la table des matières."
        )

        fld_char3 = OxmlElement("w:fldChar")
        fld_char3.set(qn("w:fldCharType"), "end")

        run._r.append(fld_char1)
        run._r.append(instr_text)
        run._r.append(fld_char2)
        run._r.append(fld_text)
        run._r.append(fld_char3)

        self.document.add_page_break()

    # ------------------------------------------------------------------
    # Chapitres et sections
    # ------------------------------------------------------------------

    def add_chapter(self, title: str, intro_text: str = "") -> None:
        """Ajoute un chapitre (Heading 1) avec un paragraphe d'introduction."""
        self.document.add_heading(title, level=1)
        if intro_text:
            para = self.document.add_paragraph(intro_text)
            para.style = self.document.styles["Normal"]

    def add_section(self, title: str, text: str = "") -> None:
        """Ajoute une section (Heading 2) avec un paragraphe de texte."""
        self.document.add_heading(title, level=2)
        if text:
            para = self.document.add_paragraph(text)
            para.style = self.document.styles["Normal"]

    def add_subsection(self, title: str, text: str = "") -> None:
        """Ajoute une sous-section (Heading 3) avec un paragraphe de texte."""
        self.document.add_heading(title, level=3)
        if text:
            para = self.document.add_paragraph(text)
            para.style = self.document.styles["Normal"]

    def add_paragraph(self, text: str) -> None:
        """Ajoute un paragraphe de corps de texte."""
        para = self.document.add_paragraph(text)
        para.style = self.document.styles["Normal"]

    def add_section_subtitle(self, text: str) -> None:
        """Ajoute un sous-titre en italique bleu ciel pour structurer les sections."""
        para = self.document.add_paragraph()
        para.paragraph_format.space_before = Pt(8)
        para.paragraph_format.space_after = Pt(4)
        run = para.add_run(text)
        run.font.name = FONT_NAME
        run.font.size = FONT_SIZE_H3
        run.font.italic = True
        run.font.bold = False
        run.font.color.rgb = COLOR_UA_BLUE

    def add_blue_arrow_paragraph(self, text: str) -> None:
        """
        Ajoute un paragraphe avec une flèche bleue ➠ en préfixe.
        Utilisé pour mettre en valeur les données clés et points importants.
        """
        para = self.document.add_paragraph()
        para.paragraph_format.space_before = Pt(4)
        para.paragraph_format.space_after = Pt(4)
        para.paragraph_format.left_indent = Cm(0.5)

        # Flèche en bleu
        arrow_run = para.add_run(f"{ARROW_LEVEL1} ")
        arrow_run.font.name = FONT_NAME
        arrow_run.font.size = FONT_SIZE_BODY
        arrow_run.font.bold = True
        arrow_run.font.color.rgb = COLOR_UA_BLUE

        # Texte en noir
        text_run = para.add_run(text)
        text_run.font.name = FONT_NAME
        text_run.font.size = FONT_SIZE_BODY
        text_run.font.color.rgb = COLOR_BLACK

    def add_nested_arrow_list(
        self,
        items: List[Union[str, tuple]],
    ) -> None:
        """
        Ajoute une liste imbriquée avec flèches bleues.

        Chaque élément peut être :
        - Une chaîne de caractères (niveau 1 : ➠)
        - Un tuple (str, list[str]) : texte niveau 1 + sous-éléments niveau 2 (➠➠)

        Exemple :
            builder.add_nested_arrow_list([
                "Point principal 1",
                ("Point principal 2", ["Sous-point A", "Sous-point B"]),
                "Point principal 3",
            ])
        """
        for item in items:
            if isinstance(item, tuple):
                label, sub_items = item
                # Niveau 1
                para = self.document.add_paragraph()
                para.paragraph_format.space_before = Pt(3)
                para.paragraph_format.space_after = Pt(3)
                para.paragraph_format.left_indent = Cm(0.5)
                arrow_run = para.add_run(f"{ARROW_LEVEL1} ")
                arrow_run.font.name = FONT_NAME
                arrow_run.font.size = FONT_SIZE_BODY
                arrow_run.font.bold = True
                arrow_run.font.color.rgb = COLOR_UA_BLUE
                text_run = para.add_run(label)
                text_run.font.name = FONT_NAME
                text_run.font.size = FONT_SIZE_BODY
                text_run.font.color.rgb = COLOR_BLACK

                # Niveau 2
                for sub in sub_items:
                    sub_para = self.document.add_paragraph()
                    sub_para.paragraph_format.space_before = Pt(2)
                    sub_para.paragraph_format.space_after = Pt(2)
                    sub_para.paragraph_format.left_indent = Cm(1.2)
                    sub_arrow = sub_para.add_run(f"{ARROW_LEVEL2} ")
                    sub_arrow.font.name = FONT_NAME
                    sub_arrow.font.size = Pt(10)
                    sub_arrow.font.color.rgb = COLOR_UA_BLUE
                    sub_text = sub_para.add_run(sub)
                    sub_text.font.name = FONT_NAME
                    sub_text.font.size = Pt(10)
                    sub_text.font.color.rgb = RGBColor(0x40, 0x40, 0x40)
            else:
                # Niveau 1 simple
                para = self.document.add_paragraph()
                para.paragraph_format.space_before = Pt(3)
                para.paragraph_format.space_after = Pt(3)
                para.paragraph_format.left_indent = Cm(0.5)
                arrow_run = para.add_run(f"{ARROW_LEVEL1} ")
                arrow_run.font.name = FONT_NAME
                arrow_run.font.size = FONT_SIZE_BODY
                arrow_run.font.bold = True
                arrow_run.font.color.rgb = COLOR_UA_BLUE
                text_run = para.add_run(item)
                text_run.font.name = FONT_NAME
                text_run.font.size = FONT_SIZE_BODY
                text_run.font.color.rgb = COLOR_BLACK

    def add_blue_box(self, text: str, border_color: str = "00A8E1") -> None:
        """
        Ajoute une boîte de texte encadrée en bleu ciel.
        Simulée par un tableau à 1 cellule avec bordure colorée et fond blanc.
        """
        table = self.document.add_table(rows=1, cols=1)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        cell = table.rows[0].cells[0]
        _set_cell_background(cell, COLOR_WHITE)
        _set_cell_border(cell, border_color)

        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        para.paragraph_format.space_before = Pt(4)
        para.paragraph_format.space_after = Pt(4)
        para.paragraph_format.left_indent = Cm(0.3)
        run = para.add_run(text)
        run.font.name = FONT_NAME
        run.font.size = FONT_SIZE_BODY
        run.font.color.rgb = COLOR_BLACK

        # Espace après la boîte
        self.document.add_paragraph()

    def add_editable_comment_zone(
        self, placeholder: str = "[À compléter par l'équipe SSU]"
    ) -> None:
        """
        Ajoute une zone de texte modifiable clairement identifiée
        pour que l'équipe SSU puisse y ajouter des commentaires.
        """
        para = self.document.add_paragraph()
        run = para.add_run(placeholder)
        run.font.name = FONT_NAME
        run.font.size = FONT_SIZE_BODY
        run.font.italic = True
        run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

    # ------------------------------------------------------------------
    # Images / graphiques
    # ------------------------------------------------------------------

    def add_image(self, image_path: str, width: float = IMAGE_WIDTH_FULL, caption: str = "") -> None:
        if not os.path.isfile(image_path):
            para = self.document.add_paragraph(f"[Image manquante : {os.path.basename(image_path)}]")
            para.runs[0].font.italic = True
            return

        # Paragraphe pour l'image
        para = self.document.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # IMPORTANT : On s'assure que l'interligne n'écrase pas l'image
        para.paragraph_format.line_spacing = 1.0 
        para.paragraph_format.space_before = Pt(12)
        para.paragraph_format.space_after = Pt(6)

        run = para.add_run()
        run.add_picture(image_path, width=width)
        
        if caption:
            cap_para = self.document.add_paragraph(caption)
            cap_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            # Empêcher la légende d'être séparée de l'image sur une autre page
            cap_para.paragraph_format.keep_with_previous = True 
            for run in cap_para.runs:
                run.font.size = FONT_SIZE_CAPTION
                run.font.italic = True
                run.font.color.rgb = RGBColor(100, 100, 100)

    # ------------------------------------------------------------------
    # Tableaux de chiffres clés
    # ------------------------------------------------------------------

    def add_key_indicators_table(
        self,
        indicators: dict,
        title: str = "Chiffres clés",
    ) -> None:
        """
        Crée un tableau à deux colonnes (Indicateur | Valeur) à partir
        d'un dictionnaire. En-tête bleu ciel, alternance gris clair / bleu pâle.
        """
        if not indicators:
            return

        if title:
            self.add_subsection(title)

        table = self.document.add_table(rows=1, cols=2)
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # En-tête
        hdr_cells = table.rows[0].cells
        for cell, text in zip(hdr_cells, ["Indicateur", "Valeur"]):
            _set_cell_background(cell, COLOR_TABLE_HEADER)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            para = cell.paragraphs[0]
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = para.add_run(text)
            run.font.name = FONT_NAME
            run.font.size = FONT_SIZE_BODY
            run.font.bold = True
            run.font.color.rgb = COLOR_WHITE

        # Lignes de données : alternance gris clair / bleu pâle
        for idx, (key, value) in enumerate(indicators.items()):
            row_cells = table.add_row().cells
            bg = COLOR_TABLE_ROW_ODD if idx % 2 == 0 else COLOR_TABLE_ROW_EVEN

            for cell in row_cells:
                _set_cell_background(cell, bg)
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

            # Nom de l'indicateur
            key_para = row_cells[0].paragraphs[0]
            key_run = key_para.add_run(str(key))
            key_run.font.name = FONT_NAME
            key_run.font.size = FONT_SIZE_BODY

            # Valeur
            val_para = row_cells[1].paragraphs[0]
            val_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            val_run = val_para.add_run(str(value))
            val_run.font.name = FONT_NAME
            val_run.font.size = FONT_SIZE_BODY
            val_run.font.bold = True

        self.document.add_paragraph()

    # ------------------------------------------------------------------
    # Sauvegarde
    # ------------------------------------------------------------------

    def save(self) -> None:
        """Sauvegarde le document Word dans le chemin spécifié."""
        os.makedirs(os.path.dirname(self.output_path) or ".", exist_ok=True)
        self.document.save(self.output_path)
        print(f"Rapport sauvegardé : {self.output_path}")
